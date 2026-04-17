import base64
import binascii
import asyncio
import hashlib
import json
import logging
import os
import secrets
import smtplib
import sqlite3
import uuid
from contextlib import closing
from datetime import datetime, timedelta
from email.utils import formataddr
from io import BytesIO
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path
from typing import Any

import httpx
from docx import Document as DocxDocument
from fastapi import Cookie, FastAPI, HTTPException, Request, Response
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pypdf import PdfReader
from pydantic import BaseModel, EmailStr, Field

from uuapi_client import (
    DEFAULT_MODEL,
    SUPPORTED_IMAGE_MEDIA_TYPES,
    SUPPORTED_MODELS,
    iter_stream_chat,
    normalize_model,
    send_chat,
)


BASE_DIR = Path(__file__).resolve().parent
DB_PATH = BASE_DIR / "chat_app.db"
DEFAULT_KEY_SOURCE_URL = "https://github.com/weicengxing/freeclaude/blob/main/key.txt"
logger = logging.getLogger(__name__)

USER_SESSION_COOKIE = "user_session"
USER_SESSION_TTL_SECONDS = 60 * 60 * 24 * 7
ROLE_USER = "User"
ROLE_SUPERADMIN = "SuperAdmin"
VALID_ROLES = {ROLE_USER, ROLE_SUPERADMIN}
VERIFY_CODE_TTL_MINUTES = 10
VERIFY_PURPOSE_REGISTER = "register"
VERIFY_PURPOSE_RESET = "reset"
DEFAULT_USER_KEY_BATCH_SIZE = 5
LEGACY_FIXED_USER_KEY_BATCH_SIZE = 5
MAX_IMAGE_BYTES = 5 * 1024 * 1024
MAX_DOCUMENT_BYTES = 10 * 1024 * 1024
STRUCTURED_MESSAGE_VERSION = 1
MAX_FILE_PROMPT_CHARS = 40_000
SUPPORTED_DOCUMENT_MEDIA_TYPES = {
    "text/plain",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
DOCUMENT_EXTENSION_MEDIA_TYPES = {
    ".txt": "text/plain",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def get_env_int(name: str, default: int, minimum: int = 1) -> int:
    raw_value = os.getenv(name)
    if raw_value is None or not raw_value.strip():
        return default
    try:
        value = int(raw_value.strip())
    except ValueError as exc:
        raise RuntimeError(f"Environment variable {name} must be an integer") from exc
    if value < minimum:
        raise RuntimeError(f"Environment variable {name} must be >= {minimum}")
    return value


USER_KEY_BATCH_SIZE = get_env_int("USER_KEY_BATCH_SIZE", DEFAULT_USER_KEY_BATCH_SIZE)
MESSAGE_RETENTION_DAYS = get_env_int("MESSAGE_RETENTION_DAYS", 14)
MESSAGE_CLEANUP_INTERVAL_DAYS = get_env_int("MESSAGE_CLEANUP_INTERVAL_DAYS", 14)
MESSAGE_CLEANUP_INTERVAL_SECONDS = MESSAGE_CLEANUP_INTERVAL_DAYS * 60 * 60 * 24

app = FastAPI(title="UUAPI Web Chat")
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))


class CreateSessionRequest(BaseModel):
    model: str = DEFAULT_MODEL


class ImagePayload(BaseModel):
    media_type: str
    data: str = Field(min_length=1)
    name: str | None = None


class FilePayload(BaseModel):
    media_type: str = ""
    data: str = Field(min_length=1)
    name: str | None = None


class ChatRequest(BaseModel):
    session_id: str
    message: str = ""
    image: ImagePayload | None = None
    images: list[ImagePayload] = Field(default_factory=list)
    files: list[FilePayload] = Field(default_factory=list)
    model: str = DEFAULT_MODEL
    replace_from_message_id: int | None = None


class UpdateSessionRequest(BaseModel):
    title: str = Field(min_length=1, max_length=200)


class EditMessageRequest(BaseModel):
    message: str = Field(min_length=1)
    model: str = DEFAULT_MODEL


class RegisterVerifyRequest(BaseModel):
    email: EmailStr


class RegisterRequest(BaseModel):
    username: str = Field(min_length=2, max_length=50)
    email: EmailStr
    password: str = Field(min_length=6, max_length=128)
    confirmPassword: str = Field(min_length=6, max_length=128)
    verifyCode: str = Field(min_length=4, max_length=12)


class LoginRequest(BaseModel):
    username: str = Field(min_length=1, max_length=100)
    password: str = Field(min_length=1, max_length=128)


class ResetVerifyRequest(BaseModel):
    email: EmailStr


class ResetPasswordRequest(BaseModel):
    email: EmailStr
    verifyCode: str = Field(min_length=4, max_length=12)
    password: str = Field(min_length=6, max_length=128)
    confirmPassword: str = Field(min_length=6, max_length=128)


class UpdateUserStatusRequest(BaseModel):
    is_active: bool


class DatabaseRowUpdateRequest(BaseModel):
    pk: dict[str, object]
    updates: dict[str, object]


class DatabaseRowInsertRequest(BaseModel):
    values: dict[str, object]


class DatabaseRowDeleteRequest(BaseModel):
    pk: dict[str, object]


def utc_now() -> str:
    return datetime.utcnow().replace(microsecond=0).isoformat() + "Z"


def now_local_text() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def get_connection() -> sqlite3.Connection:
    connection = sqlite3.connect(DB_PATH)
    connection.row_factory = sqlite3.Row
    return connection


def rollback_if_needed(connection: sqlite3.Connection) -> None:
    if connection.in_transaction:
        connection.rollback()


def table_has_column(connection: sqlite3.Connection, table_name: str, column_name: str) -> bool:
    columns = connection.execute(f"PRAGMA table_info({table_name})").fetchall()
    return any(column["name"] == column_name for column in columns)


def hash_password(password: str, salt: str) -> str:
    hashed = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt.encode("utf-8"),
        120_000,
    )
    return hashed.hex()


def generate_verify_code(length: int = 6) -> str:
    return "".join(secrets.choice("0123456789") for _ in range(length))


def env_flag(name: str, default: bool) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def get_smtp_config() -> dict[str, str | int | bool]:
    return {
        "host": os.getenv("SMTP_HOST", "smtp.qq.com").strip(),
        "port": int(os.getenv("SMTP_PORT", "587").strip()),
        "use_tls": env_flag("SMTP_USE_TLS", True),
        "username": os.getenv("SMTP_USERNAME", "").strip(),
        "password": os.getenv("SMTP_PASSWORD", "").strip(),
        "from_name": os.getenv("SMTP_FROM_NAME", "UUAPI Chat").strip() or "UUAPI Chat",
    }


def send_email(to_email: str, subject: str, html_content: str) -> bool:
    smtp_config = get_smtp_config()
    if not smtp_config["username"] or not smtp_config["password"]:
        logger.warning("SMTP is not configured; skipping email send to %s", to_email)
        return False

    try:
        msg = MIMEMultipart("alternative")
        msg["Subject"] = subject
        msg["From"] = formataddr((str(smtp_config["from_name"]), str(smtp_config["username"])))
        msg["To"] = to_email
        msg.attach(MIMEText(html_content, "html", "utf-8"))

        server = smtplib.SMTP(str(smtp_config["host"]), int(smtp_config["port"]))
        if smtp_config.get("use_tls", True):
            server.starttls()
        server.login(str(smtp_config["username"]), str(smtp_config["password"]))
        server.sendmail(str(smtp_config["username"]), [to_email], msg.as_string())
        server.quit()
        return True
    except Exception:
        logger.exception("send_email failed")
        return False


def normalize_role(role: str | None) -> str:
    role = str(role or "").strip()
    return role if role in VALID_ROLES else ROLE_USER


def model_to_dict(model: BaseModel | None) -> dict[str, Any] | None:
    if model is None:
        return None
    if hasattr(model, "model_dump"):
        return model.model_dump()
    return model.dict()


def normalize_image_payload(image: ImagePayload | dict[str, Any] | None) -> dict[str, str] | None:
    if image is None:
        return None

    raw = model_to_dict(image) if isinstance(image, BaseModel) else image
    if raw is None:
        return None

    media_type = str(raw.get("media_type", "")).strip().lower()
    data = str(raw.get("data", "")).strip()
    name = str(raw.get("name", "")).strip()

    if media_type not in SUPPORTED_IMAGE_MEDIA_TYPES:
        raise HTTPException(status_code=400, detail="Unsupported image format")
    if not data:
        raise HTTPException(status_code=400, detail="Image data cannot be empty")

    try:
        image_bytes = base64.b64decode(data, validate=True)
    except (binascii.Error, ValueError) as exc:
        raise HTTPException(status_code=400, detail="Image data must be valid base64") from exc

    if not image_bytes:
        raise HTTPException(status_code=400, detail="Image data cannot be empty")
    if len(image_bytes) > MAX_IMAGE_BYTES:
        raise HTTPException(status_code=400, detail=f"Image too large; max {MAX_IMAGE_BYTES // (1024 * 1024)} MB")

    normalized = {
        "media_type": media_type,
        "data": base64.b64encode(image_bytes).decode("ascii"),
    }
    if name:
        normalized["name"] = name[:255]
    return normalized


def normalize_image_payloads(images: list[ImagePayload | dict[str, Any]] | None) -> list[dict[str, str]]:
    normalized_images: list[dict[str, str]] = []
    for image in images or []:
        normalized_image = normalize_image_payload(image)
        if normalized_image is not None:
            normalized_images.append(normalized_image)
    return normalized_images


def normalize_request_images(
    image: ImagePayload | dict[str, Any] | None,
    images: list[ImagePayload | dict[str, Any]] | None,
) -> list[dict[str, str]]:
    if images:
        return normalize_image_payloads(images)

    normalized_image = normalize_image_payload(image)
    return [normalized_image] if normalized_image is not None else []


def normalize_document_media_type(media_type: str | None, name: str | None) -> str:
    normalized_media_type = str(media_type or "").strip().lower()
    normalized_name = str(name or "").strip().lower()
    guessed_media_type = DOCUMENT_EXTENSION_MEDIA_TYPES.get(Path(normalized_name).suffix)
    if normalized_media_type in SUPPORTED_DOCUMENT_MEDIA_TYPES:
        return normalized_media_type
    if normalized_media_type in {"", "application/octet-stream"} and guessed_media_type:
        return guessed_media_type
    return normalized_media_type


def normalize_file_payload(file_payload: FilePayload | dict[str, Any] | None) -> dict[str, str] | None:
    if file_payload is None:
        return None

    raw = model_to_dict(file_payload) if isinstance(file_payload, BaseModel) else file_payload
    if raw is None:
        return None

    name = str(raw.get("name", "")).strip()
    media_type = normalize_document_media_type(raw.get("media_type"), name)
    data = str(raw.get("data", "")).strip()

    if media_type not in SUPPORTED_DOCUMENT_MEDIA_TYPES:
        raise HTTPException(status_code=400, detail="Unsupported file format; only txt, docx, pdf are allowed")
    if not data:
        raise HTTPException(status_code=400, detail="File data cannot be empty")

    try:
        file_bytes = base64.b64decode(data, validate=True)
    except (binascii.Error, ValueError) as exc:
        raise HTTPException(status_code=400, detail="File data must be valid base64") from exc

    if not file_bytes:
        raise HTTPException(status_code=400, detail="File data cannot be empty")
    if len(file_bytes) > MAX_DOCUMENT_BYTES:
        raise HTTPException(
            status_code=400,
            detail=f"File too large; max {MAX_DOCUMENT_BYTES // (1024 * 1024)} MB",
        )

    normalized = {
        "media_type": media_type,
        "data": base64.b64encode(file_bytes).decode("ascii"),
    }
    if name:
        normalized["name"] = name[:255]
    return normalized


def normalize_request_files(files: list[FilePayload | dict[str, Any]] | None) -> list[dict[str, str]]:
    normalized_files: list[dict[str, str]] = []
    for file_payload in files or []:
        normalized_file = normalize_file_payload(file_payload)
        if normalized_file is not None:
            normalized_files.append(normalized_file)
    return normalized_files


def assign_message_images(target: dict[str, Any], images: list[dict[str, str]]) -> dict[str, Any]:
    if images:
        target["images"] = images
        target["image"] = images[0]
    return target


def assign_message_files(target: dict[str, Any], files: list[dict[str, Any]]) -> dict[str, Any]:
    if files:
        target["files"] = files
    return target


def decode_text_file_bytes(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def extract_docx_text(file_bytes: bytes) -> str:
    document = DocxDocument(BytesIO(file_bytes))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            pages.append(f"[Page {index}]\n{page_text}")
    return "\n\n".join(pages)


def parse_uploaded_file(file_payload: dict[str, str]) -> tuple[str, str]:
    media_type = str(file_payload.get("media_type", "")).strip().lower()
    file_bytes = base64.b64decode(file_payload["data"])
    if media_type == "text/plain":
        return decode_text_file_bytes(file_bytes).strip(), "text"
    if media_type == "application/pdf":
        return extract_pdf_text(file_bytes).strip(), "pypdf"
    if media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_docx_text(file_bytes).strip(), "python-docx"
    raise HTTPException(status_code=400, detail="Unsupported file format")


def prepare_message_files(files: list[dict[str, str]]) -> list[dict[str, Any]]:
    prepared_files: list[dict[str, Any]] = []
    for file_payload in files:
        parsed_text, parser_name = parse_uploaded_file(file_payload)
        file_name = str(file_payload.get("name", "")).strip() or "attachment"
        if not parsed_text:
            raise HTTPException(
                status_code=400,
                detail=f"Could not extract readable text from {file_name}. Scanned PDF/image-only files are not supported yet.",
            )
        prepared_files.append(
            {
                **file_payload,
                "name": file_name[:255],
                "parsed_text": parsed_text,
                "parser_name": parser_name,
                "parsed_char_count": len(parsed_text),
            }
        )
    return prepared_files


def build_file_prompt(files: list[dict[str, Any]] | None) -> str:
    if not files:
        return ""

    parts = [
        "以下是用户上传附件解析后的文本，请把它们当作用户补充提供的材料来理解和回答。",
        "如果附件解析结果里有格式噪声、页眉页脚或乱码，请自行甄别，不要把这些噪声当作用户真实意图。",
    ]
    for index, file_payload in enumerate(files, start=1):
        parsed_text = str(file_payload.get("parsed_text", "") or "").strip()
        if not parsed_text:
            continue
        truncated = parsed_text
        if len(truncated) > MAX_FILE_PROMPT_CHARS:
            truncated = truncated[:MAX_FILE_PROMPT_CHARS].rstrip() + "\n\n[附件解析内容过长，后续内容已截断]"
        parts.append(
            "\n".join(
                [
                    f"[附件 {index}]",
                    f"文件名: {file_payload.get('name') or f'attachment-{index}'}",
                    f"文件类型: {file_payload.get('media_type') or 'unknown'}",
                    "解析文本:",
                    truncated,
                ]
            )
        )
    return "\n\n".join(parts)


def build_model_message_content(content: str, files: list[dict[str, Any]] | None = None) -> str:
    file_prompt = build_file_prompt(files)
    text = str(content or "").strip()
    if text and file_prompt:
        return f"{text}\n\n{file_prompt}"
    if file_prompt:
        return file_prompt
    return text


def parse_stored_message_content(raw_content: str) -> dict[str, Any]:
    if raw_content:
        stripped = raw_content.lstrip()
        if stripped.startswith("{"):
            try:
                payload = json.loads(raw_content)
            except json.JSONDecodeError:
                payload = None
            if isinstance(payload, dict) and payload.get("v") == STRUCTURED_MESSAGE_VERSION:
                text = str(payload.get("text", "") or "")
                images: list[dict[str, str]] = []
                try:
                    if isinstance(payload.get("images"), list):
                        images = normalize_image_payloads(payload.get("images"))
                    elif isinstance(payload.get("image"), dict):
                        legacy_image = normalize_image_payload(payload.get("image"))
                        if legacy_image is not None:
                            images = [legacy_image]
                except HTTPException:
                    images = []
                return assign_message_images({"content": text}, images)
    return {"content": raw_content}


def serialize_message_content(content: str, images: list[dict[str, str]] | None = None) -> str:
    if not images:
        return content
    payload: dict[str, Any] = {
        "v": STRUCTURED_MESSAGE_VERSION,
        "text": content,
    }
    if len(images) == 1:
        payload["image"] = images[0]
    else:
        payload["images"] = images
    return json.dumps(payload, ensure_ascii=False, separators=(",", ":"))


def build_title_for_message(
    message_text: str,
    images: list[dict[str, Any]] | None = None,
    files: list[dict[str, Any]] | None = None,
) -> str:
    compact = " ".join(message_text.split()).strip()
    if compact:
        return build_title(compact)
    if images:
        return "图片消息"
    if files:
        return "文件消息"
    return "New Chat"


def ensure_default_admins(connection: sqlite3.Connection) -> None:
    admins = connection.execute(
        "SELECT id FROM users WHERE role = ? ORDER BY id ASC",
        (ROLE_SUPERADMIN,),
    ).fetchall()
    if len(admins) >= 2:
        return

    defaults = [
        ("admin", "admin1@example.com"),
        ("root", "admin2@example.com"),
    ]
    used_usernames = {row["username"] for row in connection.execute("SELECT username FROM users").fetchall()}
    used_emails = {row["email"] for row in connection.execute("SELECT email FROM users").fetchall()}

    for username, email in defaults:
        if len(admins) >= 2:
            break
        if username in used_usernames or email in used_emails:
            continue
        salt = secrets.token_hex(16)
        connection.execute(
            """
            INSERT INTO users (username, email, role, password_hash, salt, is_active, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                username,
                email,
                ROLE_SUPERADMIN,
                hash_password("admin123456", salt),
                salt,
                1,
                now_local_text(),
            ),
        )
        admins.append({"id": connection.execute("SELECT last_insert_rowid()").fetchone()[0]})
        used_usernames.add(username)
        used_emails.add(email)


def get_default_session_owner_id(connection: sqlite3.Connection) -> int | None:
    row = connection.execute(
        """
        SELECT id
        FROM users
        WHERE role = ? AND is_active = 1
        ORDER BY id ASC
        LIMIT 1
        """,
        (ROLE_SUPERADMIN,),
    ).fetchone()
    if row is not None:
        return int(row["id"])

    row = connection.execute(
        """
        SELECT id
        FROM users
        WHERE is_active = 1
        ORDER BY id ASC
        LIMIT 1
        """
    ).fetchone()
    return int(row["id"]) if row is not None else None


def migrate_legacy_sessions(connection: sqlite3.Connection) -> None:
    default_owner_id = get_default_session_owner_id(connection)
    if default_owner_id is None:
        return
    connection.execute(
        "UPDATE sessions SET user_id = ? WHERE user_id IS NULL",
        (default_owner_id,),
    )


def cleanup_expired_verify_codes(connection: sqlite3.Connection) -> None:
    connection.execute("DELETE FROM verify_codes WHERE expires_at <= ?", (utc_now(),))


def upsert_verify_code(connection: sqlite3.Connection, email: str, purpose: str, code: str) -> None:
    cleanup_expired_verify_codes(connection)
    now = utc_now()
    expires_at = (datetime.utcnow() + timedelta(minutes=VERIFY_CODE_TTL_MINUTES)).replace(microsecond=0).isoformat() + "Z"
    connection.execute(
        """
        INSERT INTO verify_codes (email, purpose, code, created_at, expires_at)
        VALUES (?, ?, ?, ?, ?)
        ON CONFLICT(email, purpose) DO UPDATE SET
            code = excluded.code,
            created_at = excluded.created_at,
            expires_at = excluded.expires_at
        """,
        (email, purpose, code, now, expires_at),
    )
    connection.commit()


def get_verify_code_record(connection: sqlite3.Connection, email: str, purpose: str) -> dict | None:
    cleanup_expired_verify_codes(connection)
    connection.commit()
    row = connection.execute(
        """
        SELECT email, purpose, code, created_at, expires_at
        FROM verify_codes
        WHERE email = ? AND purpose = ?
        LIMIT 1
        """,
        (email, purpose),
    ).fetchone()
    return dict(row) if row is not None else None


def delete_verify_code(connection: sqlite3.Connection, email: str, purpose: str) -> None:
    connection.execute(
        "DELETE FROM verify_codes WHERE email = ? AND purpose = ?",
        (email, purpose),
    )
    connection.commit()


def init_db() -> None:
    with closing(get_connection()) as connection:
        connection.executescript(
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT NOT NULL UNIQUE,
                email TEXT NOT NULL UNIQUE,
                role TEXT NOT NULL DEFAULT 'User',
                password_hash TEXT NOT NULL,
                salt TEXT NOT NULL,
                is_active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS auth_sessions (
                token TEXT PRIMARY KEY,
                user_id INTEGER NOT NULL,
                created_at TEXT NOT NULL,
                expires_at TEXT NOT NULL,
                FOREIGN KEY (user_id) REFERENCES users(id)
            );

            CREATE TABLE IF NOT EXISTS verify_codes (
                email TEXT NOT NULL,
                purpose TEXT NOT NULL,
                code TEXT NOT NULL,
                created_at TEXT NOT NULL,
                expires_at TEXT NOT NULL,
                PRIMARY KEY (email, purpose)
            );

            CREATE TABLE IF NOT EXISTS sessions (
                id TEXT PRIMARY KEY,
                user_id INTEGER,
                title TEXT NOT NULL,
                model TEXT NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                FOREIGN KEY (user_id) REFERENCES users(id)
            );

            CREATE TABLE IF NOT EXISTS messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                session_id TEXT NOT NULL,
                role TEXT NOT NULL,
                content TEXT NOT NULL,
                created_at TEXT NOT NULL,
                FOREIGN KEY (session_id) REFERENCES sessions(id)
            );

            CREATE TABLE IF NOT EXISTS message_files (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                message_id INTEGER NOT NULL,
                session_id TEXT NOT NULL,
                name TEXT NOT NULL,
                media_type TEXT NOT NULL,
                original_data TEXT NOT NULL,
                parsed_text TEXT NOT NULL,
                parser_name TEXT NOT NULL,
                created_at TEXT NOT NULL,
                FOREIGN KEY (message_id) REFERENCES messages(id),
                FOREIGN KEY (session_id) REFERENCES sessions(id)
            );

            CREATE TABLE IF NOT EXISTS api_keys (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                api_key TEXT NOT NULL,
                source_url TEXT NOT NULL,
                created_at TEXT NOT NULL
            );

            CREATE UNIQUE INDEX IF NOT EXISTS idx_api_keys_api_key
            ON api_keys(api_key);

            CREATE INDEX IF NOT EXISTS idx_message_files_message_id
            ON message_files(message_id);

            CREATE INDEX IF NOT EXISTS idx_message_files_session_id
            ON message_files(session_id);

            CREATE TABLE IF NOT EXISTS api_key_allocator_state (
                singleton INTEGER PRIMARY KEY CHECK (singleton = 1),
                next_batch_start_id INTEGER,
                updated_at TEXT NOT NULL
            );
            """
        )
        if not table_has_column(connection, "sessions", "user_id"):
            connection.execute("ALTER TABLE sessions ADD COLUMN user_id INTEGER")
        if not table_has_column(connection, "users", "api_key_batch_start_id"):
            connection.execute("ALTER TABLE users ADD COLUMN api_key_batch_start_id INTEGER")
        if not table_has_column(connection, "users", "current_api_key_id"):
            connection.execute("ALTER TABLE users ADD COLUMN current_api_key_id INTEGER")
        if not table_has_column(connection, "users", "api_key_batch_size"):
            connection.execute("ALTER TABLE users ADD COLUMN api_key_batch_size INTEGER")
        connection.execute(
            """
            UPDATE users
            SET api_key_batch_size = ?
            WHERE api_key_batch_size IS NULL
              AND api_key_batch_start_id IS NOT NULL
              AND current_api_key_id IS NOT NULL
            """,
            (LEGACY_FIXED_USER_KEY_BATCH_SIZE,),
        )
        ensure_default_admins(connection)
        migrate_legacy_sessions(connection)
        cleanup_expired_verify_codes(connection)
        connection.commit()


def create_auth_session(connection: sqlite3.Connection, user_id: int) -> str:
    token = secrets.token_hex(24)
    created_at = utc_now()
    expires_at = (datetime.utcnow() + timedelta(seconds=USER_SESSION_TTL_SECONDS)).replace(microsecond=0).isoformat() + "Z"
    connection.execute(
        """
        INSERT INTO auth_sessions (token, user_id, created_at, expires_at)
        VALUES (?, ?, ?, ?)
        """,
        (token, user_id, created_at, expires_at),
    )
    connection.commit()
    return token


def delete_auth_session(connection: sqlite3.Connection, token: str | None) -> None:
    if not token:
        return
    connection.execute("DELETE FROM auth_sessions WHERE token = ?", (token,))
    connection.commit()


def get_current_user(connection: sqlite3.Connection, token: str | None) -> dict | None:
    if not token:
        return None

    row = connection.execute(
        """
        SELECT
            auth_sessions.token,
            auth_sessions.expires_at,
            users.id,
            users.username,
            users.email,
            users.role,
            users.is_active
        FROM auth_sessions
        JOIN users ON users.id = auth_sessions.user_id
        WHERE auth_sessions.token = ?
        LIMIT 1
        """,
        (token,),
    ).fetchone()
    if row is None:
        return None
    if row["is_active"] != 1:
        delete_auth_session(connection, token)
        return None

    expires_at = datetime.fromisoformat(row["expires_at"].replace("Z", "+00:00"))
    if expires_at <= datetime.utcnow().astimezone(expires_at.tzinfo):
        delete_auth_session(connection, token)
        return None

    return {
        "id": row["id"],
        "username": row["username"],
        "email": row["email"],
        "role": normalize_role(row["role"]),
    }


def require_user(connection: sqlite3.Connection, token: str | None) -> dict:
    user = get_current_user(connection, token)
    if user is None:
        raise HTTPException(status_code=401, detail="请先登录")
    return user


def require_superadmin(connection: sqlite3.Connection, token: str | None) -> dict:
    user = require_user(connection, token)
    if user["role"] != ROLE_SUPERADMIN:
        raise HTTPException(status_code=403, detail="仅 SuperAdmin 可操作")
    return user


def session_exists(connection: sqlite3.Connection, session_id: str, user_id: int) -> bool:
    row = connection.execute(
        "SELECT 1 FROM sessions WHERE id = ? AND user_id = ?",
        (session_id, user_id),
    ).fetchone()
    return row is not None


def build_title(first_message: str) -> str:
    compact = " ".join(first_message.strip().split())
    return compact[:40] or "New Chat"


def create_session_record(
    connection: sqlite3.Connection,
    session_id: str,
    user_id: int,
    model: str,
    title: str = "New Chat",
) -> None:
    now = utc_now()
    connection.execute(
        """
        INSERT INTO sessions (id, user_id, title, model, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (session_id, user_id, title, normalize_model(model), now, now),
    )
    connection.commit()


def add_message(
    connection: sqlite3.Connection,
    session_id: str,
    role: str,
    content: str,
    images: list[dict[str, str]] | None = None,
    files: list[dict[str, Any]] | None = None,
    *,
    message_id: int | None = None,
    created_at: str | None = None,
) -> int:
    now = created_at or utc_now()
    if message_id is None:
        cursor = connection.execute(
            """
            INSERT INTO messages (session_id, role, content, created_at)
            VALUES (?, ?, ?, ?)
            """,
            (session_id, role, serialize_message_content(content, images), now),
        )
        inserted_message_id = int(cursor.lastrowid)
    else:
        connection.execute(
            """
            INSERT INTO messages (id, session_id, role, content, created_at)
            VALUES (?, ?, ?, ?, ?)
            """,
            (message_id, session_id, role, serialize_message_content(content, images), now),
        )
        inserted_message_id = message_id
    insert_message_files(connection, inserted_message_id, session_id, files or [], created_at=now)
    connection.execute(
        "UPDATE sessions SET updated_at = ? WHERE id = ?",
        (now, session_id),
    )
    connection.commit()
    return inserted_message_id


def insert_message_files(
    connection: sqlite3.Connection,
    message_id: int,
    session_id: str,
    files: list[dict[str, Any]],
    *,
    created_at: str,
) -> None:
    for file_payload in files:
        connection.execute(
            """
            INSERT INTO message_files (
                message_id,
                session_id,
                name,
                media_type,
                original_data,
                parsed_text,
                parser_name,
                created_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                message_id,
                session_id,
                str(file_payload.get("name", "") or "attachment")[:255],
                str(file_payload.get("media_type", "") or ""),
                str(file_payload.get("data", "") or ""),
                str(file_payload.get("parsed_text", "") or ""),
                str(file_payload.get("parser_name", "") or "unknown"),
                created_at,
            ),
        )


def build_message_file_row(row: sqlite3.Row, include_file_content: bool, include_original_data: bool) -> dict[str, Any]:
    item = {
        "id": row["id"],
        "message_id": row["message_id"],
        "session_id": row["session_id"],
        "name": row["name"],
        "media_type": row["media_type"],
        "parser_name": row["parser_name"],
        "created_at": row["created_at"],
        "parsed_char_count": int(row["parsed_char_count"] or 0),
    }
    if include_file_content:
        item["parsed_text"] = row["parsed_text"]
    if include_original_data:
        item["data"] = row["original_data"]
    return item


def get_message_files_map(
    connection: sqlite3.Connection,
    message_ids: list[int],
    *,
    include_file_content: bool = False,
    include_original_data: bool = False,
) -> dict[int, list[dict[str, Any]]]:
    if not message_ids:
        return {}

    placeholders = ",".join("?" for _ in message_ids)
    select_columns = [
        "id",
        "message_id",
        "session_id",
        "name",
        "media_type",
        "parser_name",
        "created_at",
        "LENGTH(parsed_text) AS parsed_char_count",
    ]
    if include_file_content:
        select_columns.append("parsed_text")
    if include_original_data:
        select_columns.append("original_data")

    rows = connection.execute(
        f"""
        SELECT {", ".join(select_columns)}
        FROM message_files
        WHERE message_id IN ({placeholders})
        ORDER BY id ASC
        """,
        message_ids,
    ).fetchall()
    mapping: dict[int, list[dict[str, Any]]] = {}
    for row in rows:
        message_id = int(row["message_id"])
        mapping.setdefault(message_id, []).append(
            build_message_file_row(row, include_file_content, include_original_data)
        )
    return mapping


def delete_message_files_by_ids(connection: sqlite3.Connection, message_ids: list[int]) -> None:
    if not message_ids:
        return
    placeholders = ",".join("?" for _ in message_ids)
    connection.execute(
        f"DELETE FROM message_files WHERE message_id IN ({placeholders})",
        message_ids,
    )


def delete_message_files_by_query(connection: sqlite3.Connection, query: str, params: tuple[Any, ...]) -> None:
    message_ids = [int(row["id"]) for row in connection.execute(query, params).fetchall()]
    delete_message_files_by_ids(connection, message_ids)


def update_session_metadata(
    connection: sqlite3.Connection,
    session_id: str,
    user_id: int,
    model: str,
    title: str | None = None,
) -> None:
    now = utc_now()
    if title is None:
        connection.execute(
            "UPDATE sessions SET model = ?, updated_at = ? WHERE id = ? AND user_id = ?",
            (normalize_model(model), now, session_id, user_id),
        )
    else:
        connection.execute(
            "UPDATE sessions SET title = ?, model = ?, updated_at = ? WHERE id = ? AND user_id = ?",
            (title, normalize_model(model), now, session_id, user_id),
        )
    connection.commit()


def list_sessions(connection: sqlite3.Connection, user_id: int) -> list[dict]:
    rows = connection.execute(
        """
        SELECT id, user_id, title, model, created_at, updated_at
        FROM sessions
        WHERE user_id = ?
        ORDER BY updated_at DESC, created_at DESC
        """,
        (user_id,),
    ).fetchall()
    return [dict(row) for row in rows]


def cleanup_expired_session_messages(
    connection: sqlite3.Connection,
    session_id: str,
    user_id: int,
    retention_days: int = MESSAGE_RETENTION_DAYS,
) -> None:
    cutoff = (datetime.utcnow() - timedelta(days=retention_days)).replace(microsecond=0).isoformat() + "Z"
    delete_message_files_by_query(
        connection,
        """
        SELECT messages.id
        FROM messages
        JOIN sessions ON sessions.id = messages.session_id
        WHERE messages.session_id = ?
          AND messages.created_at < ?
          AND sessions.id = ?
          AND sessions.user_id = ?
        """,
        (session_id, cutoff, session_id, user_id),
    )
    connection.execute(
        """
        DELETE FROM messages
        WHERE session_id = ?
          AND created_at < ?
          AND session_id IN (SELECT id FROM sessions WHERE id = ? AND user_id = ?)
        """,
        (session_id, cutoff, session_id, user_id),
    )
    connection.commit()


def cleanup_expired_messages(
    connection: sqlite3.Connection,
    retention_days: int = MESSAGE_RETENTION_DAYS,
) -> int:
    cutoff = (datetime.utcnow() - timedelta(days=retention_days)).replace(microsecond=0).isoformat() + "Z"
    delete_message_files_by_query(
        connection,
        "SELECT id FROM messages WHERE created_at < ?",
        (cutoff,),
    )
    cursor = connection.execute(
        """
        DELETE FROM messages
        WHERE created_at < ?
        """,
        (cutoff,),
    )
    connection.commit()
    return max(cursor.rowcount, 0)


def get_messages(
    connection: sqlite3.Connection,
    session_id: str,
    user_id: int,
    *,
    include_file_content: bool = False,
    include_original_data: bool = False,
) -> list[dict]:
    rows = connection.execute(
        """
        SELECT messages.id, messages.session_id, messages.role, messages.content, messages.created_at
        FROM messages
        JOIN sessions ON sessions.id = messages.session_id
        WHERE messages.session_id = ? AND sessions.user_id = ?
        ORDER BY messages.id ASC
        """,
        (session_id, user_id),
    ).fetchall()
    file_map = get_message_files_map(
        connection,
        [int(row["id"]) for row in rows],
        include_file_content=include_file_content,
        include_original_data=include_original_data,
    )
    messages: list[dict] = []
    for row in rows:
        item = dict(row)
        parsed = parse_stored_message_content(item["content"])
        item["content"] = parsed["content"]
        assign_message_images(item, parsed.get("images") or [])
        assign_message_files(item, file_map.get(int(item["id"]), []))
        messages.append(item)
    return messages


def get_session(connection: sqlite3.Connection, session_id: str, user_id: int) -> dict | None:
    session = connection.execute(
        """
        SELECT id, user_id, title, model, created_at, updated_at
        FROM sessions
        WHERE id = ? AND user_id = ?
        """,
        (session_id, user_id),
    ).fetchone()
    return dict(session) if session is not None else None


def get_message(
    connection: sqlite3.Connection,
    session_id: str,
    user_id: int,
    message_id: int,
    *,
    include_file_content: bool = False,
    include_original_data: bool = False,
) -> dict | None:
    row = connection.execute(
        """
        SELECT messages.id, messages.session_id, messages.role, messages.content, messages.created_at
        FROM messages
        JOIN sessions ON sessions.id = messages.session_id
        WHERE messages.session_id = ? AND sessions.user_id = ? AND messages.id = ?
        """,
        (session_id, user_id, message_id),
    ).fetchone()
    if row is None:
        return None
    item = dict(row)
    parsed = parse_stored_message_content(item["content"])
    item["content"] = parsed["content"]
    assign_message_images(item, parsed.get("images") or [])
    assign_message_files(
        item,
        get_message_files_map(
            connection,
            [int(message_id)],
            include_file_content=include_file_content,
            include_original_data=include_original_data,
        ).get(int(message_id), []),
    )
    return item


def delete_messages_from(connection: sqlite3.Connection, session_id: str, user_id: int, message_id: int) -> None:
    now = utc_now()
    delete_message_files_by_query(
        connection,
        """
        SELECT messages.id
        FROM messages
        JOIN sessions ON sessions.id = messages.session_id
        WHERE messages.session_id = ?
          AND messages.id >= ?
          AND sessions.id = ?
          AND sessions.user_id = ?
        """,
        (session_id, message_id, session_id, user_id),
    )
    connection.execute(
        """
        DELETE FROM messages
        WHERE session_id = ?
          AND id >= ?
          AND session_id IN (SELECT id FROM sessions WHERE id = ? AND user_id = ?)
        """,
        (session_id, message_id, session_id, user_id),
    )
    connection.execute(
        "UPDATE sessions SET updated_at = ? WHERE id = ? AND user_id = ?",
        (now, session_id, user_id),
    )
    connection.commit()


def delete_session_record(connection: sqlite3.Connection, session_id: str, user_id: int) -> None:
    delete_message_files_by_query(
        connection,
        """
        SELECT messages.id
        FROM messages
        JOIN sessions ON sessions.id = messages.session_id
        WHERE messages.session_id = ?
          AND sessions.id = ?
          AND sessions.user_id = ?
        """,
        (session_id, session_id, user_id),
    )
    connection.execute(
        """
        DELETE FROM messages
        WHERE session_id = ?
          AND session_id IN (SELECT id FROM sessions WHERE id = ? AND user_id = ?)
        """,
        (session_id, session_id, user_id),
    )
    connection.execute("DELETE FROM sessions WHERE id = ? AND user_id = ?", (session_id, user_id))
    connection.commit()


def trim_title(title: str) -> str:
    compact = " ".join(title.strip().split())
    return compact[:200] or "New Chat"


def normalize_github_raw_url(url: str) -> str:
    stripped = url.strip()
    if not stripped:
        raise HTTPException(status_code=400, detail="Key source URL cannot be empty")

    if stripped.startswith("https://raw.githubusercontent.com/"):
        return stripped

    marker = "https://github.com/"
    if stripped.startswith(marker) and "/blob/" in stripped:
        path = stripped[len(marker):]
        owner_repo, remainder = path.split("/blob/", 1)
        return f"https://raw.githubusercontent.com/{owner_repo}/{remainder}"

    raise HTTPException(status_code=400, detail="Unsupported GitHub URL")


def parse_api_keys_from_text(raw_text: str) -> list[str]:
    unique_keys: list[str] = []
    seen: set[str] = set()
    for line in raw_text.splitlines():
        key = line.strip()
        if not key or key in seen:
            continue
        seen.add(key)
        unique_keys.append(key)
    return unique_keys


def import_api_keys_from_url(connection: sqlite3.Connection, source_url: str) -> dict:
    raw_url = normalize_github_raw_url(source_url)
    try:
        response = httpx.get(raw_url, timeout=30.0, follow_redirects=True)
        response.raise_for_status()
    except httpx.HTTPError as exc:
        raise HTTPException(status_code=502, detail=f"Failed to fetch key file: {exc}") from exc

    parsed_keys = parse_api_keys_from_text(response.text)
    if not parsed_keys:
        raise HTTPException(status_code=400, detail="No valid keys found in the source file")

    inserted_count = 0
    now = utc_now()
    for api_key in parsed_keys:
        cursor = connection.execute(
            """
            INSERT OR IGNORE INTO api_keys (api_key, source_url, created_at)
            VALUES (?, ?, ?)
            """,
            (api_key, source_url, now),
        )
        inserted_count += cursor.rowcount
    connection.commit()

    total_keys = connection.execute("SELECT COUNT(*) AS count FROM api_keys").fetchone()["count"]
    return {
        "source_url": source_url,
        "raw_url": raw_url,
        "read_count": len(parsed_keys),
        "inserted_count": inserted_count,
        "ignored_count": len(parsed_keys) - inserted_count,
        "total_keys": total_keys,
    }


def get_api_key_id_bounds(connection: sqlite3.Connection) -> tuple[int | None, int | None]:
    row = connection.execute(
        "SELECT MIN(id) AS min_id, MAX(id) AS max_id FROM api_keys"
    ).fetchone()
    if row is None or row["min_id"] is None or row["max_id"] is None:
        return None, None
    return int(row["min_id"]), int(row["max_id"])


def get_api_key_count(connection: sqlite3.Connection) -> int:
    row = connection.execute("SELECT COUNT(*) AS count FROM api_keys").fetchone()
    return int(row["count"]) if row is not None else 0


def get_api_key_record_by_id(connection: sqlite3.Connection, key_id: int | None) -> dict | None:
    if key_id is None:
        return None
    row = connection.execute(
        """
        SELECT id, api_key, source_url, created_at
        FROM api_keys
        WHERE id = ?
        LIMIT 1
        """,
        (key_id,),
    ).fetchone()
    return dict(row) if row is not None else None


def get_user_key_state(connection: sqlite3.Connection, user_id: int) -> dict | None:
    row = connection.execute(
        """
        SELECT id, api_key_batch_start_id, current_api_key_id, api_key_batch_size
        FROM users
        WHERE id = ?
        LIMIT 1
        """,
        (user_id,),
    ).fetchone()
    return dict(row) if row is not None else None


def get_effective_user_batch_size(user_state: dict | None) -> int | None:
    if user_state is None:
        return None
    batch_size = user_state.get("api_key_batch_size")
    if batch_size is None:
        return None
    return int(batch_size)


def init_allocator_state_row(connection: sqlite3.Connection) -> None:
    row = connection.execute(
        "SELECT singleton FROM api_key_allocator_state WHERE singleton = 1"
    ).fetchone()
    if row is None:
        connection.execute(
            """
            INSERT INTO api_key_allocator_state (singleton, next_batch_start_id, updated_at)
            VALUES (1, NULL, ?)
            """,
            (utc_now(),),
        )


def allocate_key_batch_locked(connection: sqlite3.Connection, user_id: int) -> dict:
    min_key_id, _ = get_api_key_id_bounds(connection)
    if min_key_id is None:
        raise HTTPException(status_code=503, detail="当前没有可用的 API Key")

    try:
        connection.execute("BEGIN IMMEDIATE")
        init_allocator_state_row(connection)
        state = connection.execute(
            """
            SELECT next_batch_start_id
            FROM api_key_allocator_state
            WHERE singleton = 1
            LIMIT 1
            """
        ).fetchone()

        batch_start_id = min_key_id
        if state is not None and state["next_batch_start_id"] is not None:
            batch_start_id = max(int(state["next_batch_start_id"]), min_key_id)

        allocated_batch_size = USER_KEY_BATCH_SIZE
        batch_end_id = batch_start_id + allocated_batch_size - 1
        available_count = connection.execute(
            """
            SELECT COUNT(*) AS count
            FROM api_keys
            WHERE id BETWEEN ? AND ?
            """,
            (batch_start_id, batch_end_id),
        ).fetchone()
        if available_count is None or int(available_count["count"]) != allocated_batch_size:
            raise HTTPException(
                status_code=503,
                detail=f"可分配的连续 API Key 不足 {allocated_batch_size} 个",
            )

        now = utc_now()
        connection.execute(
            """
            UPDATE api_key_allocator_state
            SET next_batch_start_id = ?, updated_at = ?
            WHERE singleton = 1
            """,
            (batch_end_id + 1, now),
        )
        connection.execute(
            """
            UPDATE users
            SET api_key_batch_start_id = ?, current_api_key_id = ?, api_key_batch_size = ?
            WHERE id = ?
            """,
            (batch_start_id, batch_start_id, allocated_batch_size, user_id),
        )
        connection.commit()
    except Exception:
        rollback_if_needed(connection)
        raise

    key_record = get_api_key_record_by_id(connection, batch_start_id)
    if key_record is None:
        raise HTTPException(status_code=503, detail="分配的 API Key 不存在")
    return key_record


def ensure_user_api_key(connection: sqlite3.Connection, user_id: int) -> dict:
    user_state = get_user_key_state(connection, user_id)
    current_key_id = None if user_state is None else user_state["current_api_key_id"]
    key_record = get_api_key_record_by_id(connection, current_key_id)
    if key_record is not None:
        return key_record
    return allocate_key_batch_locked(connection, user_id)


def advance_user_api_key(connection: sqlite3.Connection, user_id: int, exhausted_key_id: int | None) -> dict:
    try:
        connection.execute("BEGIN IMMEDIATE")
        user_state = get_user_key_state(connection, user_id)
        if user_state is None:
            raise HTTPException(status_code=404, detail="用户不存在")

        current_key_id = user_state["current_api_key_id"]
        batch_start_id = user_state["api_key_batch_start_id"]
        batch_size = get_effective_user_batch_size(user_state)

        if current_key_id is None or batch_start_id is None or batch_size is None:
            rollback_if_needed(connection)
            return allocate_key_batch_locked(connection, user_id)

        if exhausted_key_id is not None and current_key_id != exhausted_key_id:
            connection.commit()
            key_record = get_api_key_record_by_id(connection, current_key_id)
            if key_record is None:
                return allocate_key_batch_locked(connection, user_id)
            return key_record

        batch_end_id = int(batch_start_id) + batch_size - 1
        if int(current_key_id) < batch_end_id:
            next_key_id = int(current_key_id) + 1
            key_record = get_api_key_record_by_id(connection, next_key_id)
            if key_record is None:
                raise HTTPException(status_code=503, detail="下一个 API Key 不存在")
            connection.execute(
                "UPDATE users SET current_api_key_id = ? WHERE id = ?",
                (next_key_id, user_id),
            )
            connection.commit()
            return key_record

        connection.commit()
        return allocate_key_batch_locked(connection, user_id)
    except Exception:
        rollback_if_needed(connection)
        raise


def is_api_key_quota_error(exc: Exception) -> bool:
    if isinstance(exc, HTTPException):
        return False

    message_parts = [str(exc)]
    if isinstance(exc, httpx.HTTPStatusError):
        message_parts.append(exc.response.text)

    normalized = " ".join(part.lower() for part in message_parts if part)
    keywords = (
        "quota",
        "余额",
        "额度",
        "insufficient",
        "credit",
        "rate limit",
        "api key has been disabled",
        "invalid x-api-key",
        "unauthorized",
        "forbidden",
    )
    return any(keyword in normalized for keyword in keywords)


def get_exception_detail(exc: Exception) -> str:
    if isinstance(exc, HTTPException):
        return str(exc.detail)
    return str(exc)


def send_chat_with_user_api_key(
    connection: sqlite3.Connection,
    user_id: int,
    messages: list[dict[str, Any]],
    model: str,
    session_id: str,
) -> dict:
    retry_limit = max(1, min(get_api_key_count(connection), 50))
    attempted_key_ids: set[int] = set()
    last_exc: Exception | None = None

    for _ in range(retry_limit):
        key_record = ensure_user_api_key(connection, user_id)
        key_id = int(key_record["id"])
        if key_id in attempted_key_ids:
            break
        attempted_key_ids.add(key_id)

        try:
            return send_chat(
                messages=messages,
                model=model,
                session_id=session_id,
                api_key=key_record["api_key"],
            )
        except Exception as exc:
            last_exc = exc
            if not is_api_key_quota_error(exc):
                raise
            advance_user_api_key(connection, user_id, key_id)

    if last_exc is not None:
        raise HTTPException(status_code=503, detail=f"所有可切换的 API Key 都不可用: {last_exc}") from last_exc
    raise HTTPException(status_code=503, detail="当前没有可用的 API Key")


def build_request_messages(history: list[dict]) -> list[dict[str, Any]]:
    request_messages: list[dict[str, Any]] = []
    for item in history:
        files = item.get("files") if isinstance(item.get("files"), list) else []
        message: dict[str, Any] = {
            "role": item["role"],
            "content": build_model_message_content(item.get("content", ""), files if item.get("role") == "user" else []),
        }
        images = item.get("images")
        if isinstance(images, list) and images:
            message["images"] = images
        elif item.get("image") is not None:
            message["images"] = [item["image"]]
        request_messages.append(message)
    return request_messages


def restore_messages(connection: sqlite3.Connection, messages: list[dict]) -> None:
    for item in messages:
        add_message(
            connection=connection,
            session_id=item["session_id"],
            role=item["role"],
            content=item.get("content", ""),
            images=item.get("images") or [],
            files=item.get("files") or [],
            message_id=int(item["id"]),
            created_at=item["created_at"],
        )
    connection.commit()


def auth_payload(user: dict | None) -> dict:
    if user is None:
        return {"authenticated": False, "user": None, "is_admin": False}
    return {
        "authenticated": True,
        "is_admin": user["role"] == ROLE_SUPERADMIN,
        "user": user,
    }


def list_users(connection: sqlite3.Connection) -> list[dict]:
    rows = connection.execute(
        """
        SELECT id, username, email, role, is_active, created_at
        FROM users
        ORDER BY role DESC, id ASC
        """
    ).fetchall()
    return [dict(row) for row in rows]


def quote_identifier(identifier: str) -> str:
    return '"' + identifier.replace('"', '""') + '"'


def get_manageable_table_names(connection: sqlite3.Connection) -> list[str]:
    rows = connection.execute(
        """
        SELECT name
        FROM sqlite_master
        WHERE type = 'table' AND name NOT LIKE 'sqlite_%'
        ORDER BY name
        """
    ).fetchall()
    return [str(row["name"]) for row in rows]


def ensure_manageable_table(connection: sqlite3.Connection, table_name: str) -> str:
    normalized = str(table_name).strip()
    if normalized not in get_manageable_table_names(connection):
        raise HTTPException(status_code=404, detail="Table not found")
    return normalized


def get_table_columns(connection: sqlite3.Connection, table_name: str) -> list[dict]:
    table_name = ensure_manageable_table(connection, table_name)
    rows = connection.execute(f"PRAGMA table_info({quote_identifier(table_name)})").fetchall()
    return [dict(row) for row in rows]


def get_table_primary_key_columns(connection: sqlite3.Connection, table_name: str) -> list[dict]:
    columns = get_table_columns(connection, table_name)
    return sorted((column for column in columns if int(column["pk"]) > 0), key=lambda item: int(item["pk"]))


def get_table_row_count(connection: sqlite3.Connection, table_name: str) -> int:
    table_name = ensure_manageable_table(connection, table_name)
    row = connection.execute(
        f"SELECT COUNT(*) AS count FROM {quote_identifier(table_name)}"
    ).fetchone()
    return int(row["count"]) if row is not None else 0


def list_tables_with_metadata(connection: sqlite3.Connection) -> list[dict]:
    results: list[dict] = []
    for table_name in get_manageable_table_names(connection):
        columns = get_table_columns(connection, table_name)
        primary_keys = [column["name"] for column in get_table_primary_key_columns(connection, table_name)]
        results.append(
            {
                "name": table_name,
                "count": get_table_row_count(connection, table_name),
                "columns": [column["name"] for column in columns],
                "primary_keys": primary_keys,
            }
        )
    return results


def list_table_rows(connection: sqlite3.Connection, table_name: str, limit: int = 50, offset: int = 0) -> dict:
    normalized_table = ensure_manageable_table(connection, table_name)
    columns = get_table_columns(connection, normalized_table)
    primary_keys = [column["name"] for column in get_table_primary_key_columns(connection, normalized_table)]
    safe_limit = max(1, min(int(limit), 200))
    safe_offset = max(0, int(offset))
    rows = connection.execute(
        f"SELECT * FROM {quote_identifier(normalized_table)} LIMIT ? OFFSET ?",
        (safe_limit, safe_offset),
    ).fetchall()
    return {
        "table": normalized_table,
        "columns": columns,
        "primary_keys": primary_keys,
        "count": get_table_row_count(connection, normalized_table),
        "rows": [dict(row) for row in rows],
        "limit": safe_limit,
        "offset": safe_offset,
    }


def validate_row_payload_columns(connection: sqlite3.Connection, table_name: str, payload: dict[str, object]) -> None:
    valid_columns = {column["name"] for column in get_table_columns(connection, table_name)}
    invalid_columns = sorted(set(payload) - valid_columns)
    if invalid_columns:
        raise HTTPException(status_code=400, detail=f"Unknown columns: {', '.join(invalid_columns)}")


def build_where_clause_from_pk(primary_key_values: dict[str, object]) -> tuple[str, list[object]]:
    if not primary_key_values:
        raise HTTPException(status_code=400, detail="Primary key values are required")
    parts: list[str] = []
    params: list[object] = []
    for key, value in primary_key_values.items():
        parts.append(f"{quote_identifier(key)} = ?")
        params.append(value)
    return " AND ".join(parts), params


def update_table_row(connection: sqlite3.Connection, table_name: str, pk: dict[str, object], updates: dict[str, object]) -> dict:
    normalized_table = ensure_manageable_table(connection, table_name)
    primary_keys = [column["name"] for column in get_table_primary_key_columns(connection, normalized_table)]
    if not primary_keys:
        raise HTTPException(status_code=400, detail="This table cannot be updated without a primary key")
    if sorted(pk.keys()) != sorted(primary_keys):
        raise HTTPException(status_code=400, detail=f"Primary key must contain exactly: {', '.join(primary_keys)}")
    if not updates:
        raise HTTPException(status_code=400, detail="No updates provided")

    validate_row_payload_columns(connection, normalized_table, pk)
    validate_row_payload_columns(connection, normalized_table, updates)
    invalid_updates = sorted(set(updates) & set(primary_keys))
    if invalid_updates:
        raise HTTPException(status_code=400, detail=f"Primary key columns cannot be updated: {', '.join(invalid_updates)}")

    set_clause = ", ".join(f"{quote_identifier(column)} = ?" for column in updates)
    set_params = list(updates.values())
    where_clause, where_params = build_where_clause_from_pk(pk)
    try:
        cursor = connection.execute(
            f"UPDATE {quote_identifier(normalized_table)} SET {set_clause} WHERE {where_clause}",
            [*set_params, *where_params],
        )
    except sqlite3.IntegrityError as exc:
        raise HTTPException(status_code=409, detail=f"Update failed: {exc}") from exc
    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="Row not found")
    connection.commit()
    row = connection.execute(
        f"SELECT * FROM {quote_identifier(normalized_table)} WHERE {where_clause} LIMIT 1",
        where_params,
    ).fetchone()
    return dict(row) if row is not None else {}


def insert_table_row(connection: sqlite3.Connection, table_name: str, values: dict[str, object]) -> dict:
    normalized_table = ensure_manageable_table(connection, table_name)
    if not values:
        raise HTTPException(status_code=400, detail="No values provided")
    validate_row_payload_columns(connection, normalized_table, values)
    columns = list(values.keys())
    placeholders = ", ".join("?" for _ in columns)
    column_clause = ", ".join(quote_identifier(column) for column in columns)
    params = [values[column] for column in columns]
    try:
        cursor = connection.execute(
            f"INSERT INTO {quote_identifier(normalized_table)} ({column_clause}) VALUES ({placeholders})",
            params,
        )
    except sqlite3.IntegrityError as exc:
        raise HTTPException(status_code=409, detail=f"Insert failed: {exc}") from exc
    connection.commit()
    primary_keys = [column["name"] for column in get_table_primary_key_columns(connection, normalized_table)]
    if len(primary_keys) == 1 and primary_keys[0] not in values:
        inserted_pk = {primary_keys[0]: cursor.lastrowid}
        where_clause, where_params = build_where_clause_from_pk(inserted_pk)
        row = connection.execute(
            f"SELECT * FROM {quote_identifier(normalized_table)} WHERE {where_clause} LIMIT 1",
            where_params,
        ).fetchone()
        if row is not None:
            return dict(row)
    return dict(values)


def delete_table_row(connection: sqlite3.Connection, table_name: str, pk: dict[str, object]) -> None:
    normalized_table = ensure_manageable_table(connection, table_name)
    primary_keys = [column["name"] for column in get_table_primary_key_columns(connection, normalized_table)]
    if not primary_keys:
        raise HTTPException(status_code=400, detail="This table cannot be deleted without a primary key")
    if sorted(pk.keys()) != sorted(primary_keys):
        raise HTTPException(status_code=400, detail=f"Primary key must contain exactly: {', '.join(primary_keys)}")
    validate_row_payload_columns(connection, normalized_table, pk)
    where_clause, where_params = build_where_clause_from_pk(pk)
    cursor = connection.execute(
        f"DELETE FROM {quote_identifier(normalized_table)} WHERE {where_clause}",
        where_params,
    )
    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="Row not found")
    connection.commit()


def send_and_persist_reply(
    connection: sqlite3.Connection,
    session_id: str,
    user_id: int,
    message_text: str,
    images: list[dict[str, str]],
    files: list[dict[str, Any]],
    model: str,
    replace_from_message_id: int | None = None,
) -> tuple[dict, list[dict], str]:
    if not session_exists(connection, session_id, user_id):
        create_session_record(connection, session_id, user_id, model)

    current_session = get_session(connection, session_id, user_id)
    removed_suffix: list[dict] = []
    if replace_from_message_id is not None:
        target_message = get_message(connection, session_id, user_id, replace_from_message_id)
        if target_message is None:
            raise HTTPException(status_code=404, detail="Message not found")
        removed_suffix = [
            item
            for item in get_messages(connection, session_id, user_id, include_file_content=True, include_original_data=True)
            if item["id"] >= replace_from_message_id
        ]
        delete_messages_from(connection, session_id, user_id, replace_from_message_id)

    history_before = get_messages(connection, session_id, user_id, include_file_content=True)
    add_message(connection, session_id, "user", message_text, images=images, files=files)

    should_autobuild_title = not history_before and (current_session is None or current_session["title"] == "New Chat")
    title = build_title_for_message(message_text, images, files) if should_autobuild_title else None
    request_messages = build_request_messages(
        [
            *history_before,
            {"role": "user", "content": message_text, "images": images, "files": files},
        ]
    )

    try:
        response = send_chat_with_user_api_key(
            connection=connection,
            user_id=user_id,
            messages=request_messages,
            model=model,
            session_id=session_id,
        )
    except HTTPException:
        latest_user = connection.execute(
            """
            SELECT id
            FROM messages
            WHERE session_id = ? AND role = 'user'
            ORDER BY id DESC
            LIMIT 1
            """,
            (session_id,),
        ).fetchone()
        if latest_user is not None:
            delete_message_files_by_ids(connection, [int(latest_user["id"])])
            connection.execute("DELETE FROM messages WHERE id = ?", (latest_user["id"],))
            connection.commit()
        if removed_suffix:
            restore_messages(connection, removed_suffix)
        raise
    except Exception as exc:
        latest_user = connection.execute(
            """
            SELECT id
            FROM messages
            WHERE session_id = ? AND role = 'user'
            ORDER BY id DESC
            LIMIT 1
            """,
            (session_id,),
        ).fetchone()
        if latest_user is not None:
            delete_message_files_by_ids(connection, [int(latest_user["id"])])
            connection.execute("DELETE FROM messages WHERE id = ?", (latest_user["id"],))
            connection.commit()
        if removed_suffix:
            restore_messages(connection, removed_suffix)
        raise HTTPException(status_code=502, detail=f"Upstream request failed: {exc}") from exc

    assistant_text = response["text"].strip() or "(empty response)"
    add_message(connection, session_id, "assistant", assistant_text)
    update_session_metadata(connection, session_id, user_id, response["model"], title=title)
    session = get_session(connection, session_id, user_id)
    messages = get_messages(connection, session_id, user_id)
    return session, messages, assistant_text


def sse_event(event: str, data: dict) -> str:
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


def extract_stream_text(chunk: dict) -> str:
    data = chunk.get("data", {})
    if data.get("type") == "content_block_delta":
        delta = data.get("delta", {})
        if delta.get("type") == "text_delta":
            return delta.get("text", "")
    return ""


async def periodic_message_cleanup() -> None:
    while True:
        try:
            with closing(get_connection()) as connection:
                deleted_count = cleanup_expired_messages(connection)
            logger.info("Periodic message cleanup finished, deleted %s expired rows", deleted_count)
        except asyncio.CancelledError:
            raise
        except Exception:
            logger.exception("Periodic message cleanup failed")
        await asyncio.sleep(MESSAGE_CLEANUP_INTERVAL_SECONDS)


@app.on_event("startup")
async def startup() -> None:
    init_db()
    app.state.message_cleanup_task = asyncio.create_task(periodic_message_cleanup())


@app.on_event("shutdown")
async def shutdown() -> None:
    task = getattr(app.state, "message_cleanup_task", None)
    if task is None:
        return
    task.cancel()
    try:
        await task
    except asyncio.CancelledError:
        pass


@app.get("/", response_class=HTMLResponse)
def index(request: Request) -> HTMLResponse:
    return templates.TemplateResponse(
        request=request,
        name="index.html",
        context={
            "models": sorted(SUPPORTED_MODELS),
            "default_model": DEFAULT_MODEL,
            "default_admin_accounts": [
                {"username": "admin", "password": "admin123456"},
                {"username": "root", "password": "admin123456"},
            ],
        },
    )


@app.get("/api/models")
def api_models() -> dict:
    return {
        "models": sorted(SUPPORTED_MODELS),
        "default_model": DEFAULT_MODEL,
    }


@app.get("/api/keys")
def api_keys_status(user_session: str | None = Cookie(default=None)) -> dict:
    with closing(get_connection()) as connection:
        require_superadmin(connection, user_session)
        total_keys = connection.execute("SELECT COUNT(*) AS count FROM api_keys").fetchone()["count"]
    return {
        "source_url": DEFAULT_KEY_SOURCE_URL,
        "total_keys": total_keys,
    }


@app.post("/api/keys/import")
def api_import_keys(user_session: str | None = Cookie(default=None)) -> dict:
    with closing(get_connection()) as connection:
        require_superadmin(connection, user_session)
        return import_api_keys_from_url(connection, DEFAULT_KEY_SOURCE_URL)


@app.get("/api/auth/me")
def api_auth_me(user_session: str | None = Cookie(default=None)) -> dict:
    with closing(get_connection()) as connection:
        user = get_current_user(connection, user_session)
        return auth_payload(user)


@app.post("/api/auth/register-verify")
def api_register_verify(payload: RegisterVerifyRequest) -> dict:
    email = payload.email.strip().lower()
    code = generate_verify_code()
    with closing(get_connection()) as connection:
        upsert_verify_code(connection, email, VERIFY_PURPOSE_REGISTER, code)

    ok = send_email(
        email,
        "注册验证码",
        f"<p>你的注册验证码是：<strong>{code}</strong></p><p>{VERIFY_CODE_TTL_MINUTES}分钟内有效。</p>",
    )
    if not ok:
        with closing(get_connection()) as connection:
            delete_verify_code(connection, email, VERIFY_PURPOSE_REGISTER)
        raise HTTPException(status_code=500, detail="验证码发送失败，请先完成 SMTP 配置")
    return {"message": "验证码已发送"}


@app.post("/api/auth/register", status_code=201)
def api_register(payload: RegisterRequest) -> dict:
    username = payload.username.strip()
    email = payload.email.strip().lower()
    password = payload.password.strip()
    confirm_password = payload.confirmPassword.strip()
    verify_code = payload.verifyCode.strip()

    if password != confirm_password:
        raise HTTPException(status_code=400, detail="两次密码不一致")

    with closing(get_connection()) as connection:
        stored = get_verify_code_record(connection, email, VERIFY_PURPOSE_REGISTER)
    current_utc = datetime.utcnow()
    if not stored:
        raise HTTPException(status_code=400, detail="请先获取验证码")
    if datetime.fromisoformat(stored["expires_at"].replace("Z", "+00:00")).replace(tzinfo=None) <= current_utc:
        with closing(get_connection()) as connection:
            delete_verify_code(connection, email, VERIFY_PURPOSE_REGISTER)
        raise HTTPException(status_code=400, detail="验证码已过期")
    if stored["code"] != verify_code:
        raise HTTPException(status_code=400, detail="验证码错误")

    with closing(get_connection()) as connection:
        exists = connection.execute(
            "SELECT id FROM users WHERE username = ? OR email = ?",
            (username, email),
        ).fetchone()
        if exists:
            raise HTTPException(status_code=409, detail="用户名或邮箱已存在")

        salt = secrets.token_hex(16)
        connection.execute(
            """
            INSERT INTO users (username, email, role, password_hash, salt, is_active, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                username,
                email,
                ROLE_USER,
                hash_password(password, salt),
                salt,
                1,
                now_local_text(),
            ),
        )
        user_id = int(connection.execute("SELECT last_insert_rowid()").fetchone()[0])
        connection.commit()
        try:
            ensure_user_api_key(connection, user_id)
        except HTTPException:
            logger.warning("register user %s without preallocated api key batch", user_id)

    with closing(get_connection()) as connection:
        delete_verify_code(connection, email, VERIFY_PURPOSE_REGISTER)
    return {"message": "注册成功", "username": username}


@app.post("/api/auth/login")
def api_login(payload: LoginRequest, response: Response) -> dict:
    username = payload.username.strip()
    password = payload.password.strip()

    with closing(get_connection()) as connection:
        row = connection.execute(
            """
            SELECT id, username, email, role, password_hash, salt, is_active
            FROM users
            WHERE username = ? OR email = ?
            LIMIT 1
            """,
            (username, username),
        ).fetchone()

        if row is None or hash_password(password, row["salt"]) != row["password_hash"]:
            raise HTTPException(status_code=401, detail="账号或密码错误")
        if row["is_active"] != 1:
            raise HTTPException(status_code=403, detail="账号已禁用")

        token = create_auth_session(connection, row["id"])

    response.set_cookie(
        key=USER_SESSION_COOKIE,
        value=token,
        httponly=True,
        max_age=USER_SESSION_TTL_SECONDS,
        path="/",
        samesite="lax",
    )
    return {
        "message": "登录成功",
        "user": {
            "id": row["id"],
            "username": row["username"],
            "email": row["email"],
            "role": normalize_role(row["role"]),
        },
    }


@app.post("/api/auth/logout")
def api_logout(response: Response, user_session: str | None = Cookie(default=None)) -> dict:
    with closing(get_connection()) as connection:
        delete_auth_session(connection, user_session)
    response.delete_cookie(key=USER_SESSION_COOKIE, path="/")
    return {"message": "已退出登录"}


@app.post("/api/auth/verify")
def api_reset_verify(payload: ResetVerifyRequest) -> dict:
    email = payload.email.strip().lower()

    with closing(get_connection()) as connection:
        user = connection.execute("SELECT id FROM users WHERE email = ?", (email,)).fetchone()
        if user is None:
            raise HTTPException(status_code=400, detail="该邮箱未注册")

    code = generate_verify_code()
    with closing(get_connection()) as connection:
        upsert_verify_code(connection, email, VERIFY_PURPOSE_RESET, code)

    ok = send_email(
        email,
        "找回密码验证码",
        f"<p>你的找回密码验证码是：<strong>{code}</strong></p><p>{VERIFY_CODE_TTL_MINUTES}分钟内有效。</p>",
    )
    if not ok:
        with closing(get_connection()) as connection:
            delete_verify_code(connection, email, VERIFY_PURPOSE_RESET)
        raise HTTPException(status_code=500, detail="验证码发送失败，请先完成 SMTP 配置")
    return {"message": "验证码已发送"}


@app.post("/api/auth/reset")
def api_reset_password(payload: ResetPasswordRequest) -> dict:
    email = payload.email.strip().lower()
    verify_code = payload.verifyCode.strip()
    password = payload.password.strip()
    confirm_password = payload.confirmPassword.strip()

    if password != confirm_password:
        raise HTTPException(status_code=400, detail="两次密码不一致")

    with closing(get_connection()) as connection:
        stored = get_verify_code_record(connection, email, VERIFY_PURPOSE_RESET)
    current_utc = datetime.utcnow()
    if not stored:
        raise HTTPException(status_code=400, detail="请先获取验证码")
    if datetime.fromisoformat(stored["expires_at"].replace("Z", "+00:00")).replace(tzinfo=None) <= current_utc:
        with closing(get_connection()) as connection:
            delete_verify_code(connection, email, VERIFY_PURPOSE_RESET)
        raise HTTPException(status_code=400, detail="验证码已过期")
    if stored["code"] != verify_code:
        raise HTTPException(status_code=400, detail="验证码错误")

    salt = secrets.token_hex(16)
    with closing(get_connection()) as connection:
        connection.execute(
            "UPDATE users SET password_hash = ?, salt = ? WHERE email = ?",
            (hash_password(password, salt), salt, email),
        )
        connection.commit()

    with closing(get_connection()) as connection:
        delete_verify_code(connection, email, VERIFY_PURPOSE_RESET)
    return {"message": "密码重置成功"}


@app.get("/api/admin/users")
def api_admin_list_users(user_session: str | None = Cookie(default=None)) -> dict:
    with closing(get_connection()) as connection:
        require_superadmin(connection, user_session)
        return {"users": list_users(connection)}


@app.patch("/api/admin/users/{user_id}")
def api_admin_update_user_status(
    user_id: int,
    payload: UpdateUserStatusRequest,
    user_session: str | None = Cookie(default=None),
) -> dict:
    with closing(get_connection()) as connection:
        admin = require_superadmin(connection, user_session)
        target = connection.execute(
            """
            SELECT id, username, email, role, is_active, created_at
            FROM users
            WHERE id = ?
            LIMIT 1
            """,
            (user_id,),
        ).fetchone()
        if target is None:
            raise HTTPException(status_code=404, detail="用户不存在")
        if target["id"] == admin["id"] and not payload.is_active:
            raise HTTPException(status_code=400, detail="不能禁用当前登录的管理员")

        connection.execute(
            "UPDATE users SET is_active = ? WHERE id = ?",
            (1 if payload.is_active else 0, user_id),
        )
        if not payload.is_active:
            connection.execute("DELETE FROM auth_sessions WHERE user_id = ?", (user_id,))
        connection.commit()

        updated = connection.execute(
            """
            SELECT id, username, email, role, is_active, created_at
            FROM users
            WHERE id = ?
            LIMIT 1
            """,
            (user_id,),
        ).fetchone()
        return {"user": dict(updated)}


@app.get("/api/admin/db/tables")
def api_admin_db_tables(user_session: str | None = Cookie(default=None)) -> dict:
    with closing(get_connection()) as connection:
        require_superadmin(connection, user_session)
        return {"tables": list_tables_with_metadata(connection)}


@app.get("/api/admin/db/tables/{table_name}")
def api_admin_db_table_rows(
    table_name: str,
    limit: int = 50,
    offset: int = 0,
    user_session: str | None = Cookie(default=None),
) -> dict:
    with closing(get_connection()) as connection:
        require_superadmin(connection, user_session)
        return list_table_rows(connection, table_name, limit=limit, offset=offset)


@app.post("/api/admin/db/tables/{table_name}/rows")
def api_admin_db_insert_row(
    table_name: str,
    payload: DatabaseRowInsertRequest,
    user_session: str | None = Cookie(default=None),
) -> dict:
    with closing(get_connection()) as connection:
        require_superadmin(connection, user_session)
        row = insert_table_row(connection, table_name, payload.values)
        return {"row": row}


@app.patch("/api/admin/db/tables/{table_name}/rows")
def api_admin_db_update_row(
    table_name: str,
    payload: DatabaseRowUpdateRequest,
    user_session: str | None = Cookie(default=None),
) -> dict:
    with closing(get_connection()) as connection:
        require_superadmin(connection, user_session)
        row = update_table_row(connection, table_name, payload.pk, payload.updates)
        return {"row": row}


@app.delete("/api/admin/db/tables/{table_name}/rows")
def api_admin_db_delete_row(
    table_name: str,
    payload: DatabaseRowDeleteRequest,
    user_session: str | None = Cookie(default=None),
) -> dict:
    with closing(get_connection()) as connection:
        require_superadmin(connection, user_session)
        delete_table_row(connection, table_name, payload.pk)
        return {"ok": True}


@app.get("/api/sessions")
def api_list_sessions(user_session: str | None = Cookie(default=None)) -> list[dict]:
    with closing(get_connection()) as connection:
        user = require_user(connection, user_session)
        return list_sessions(connection, user["id"])


@app.post("/api/sessions")
def api_create_session(payload: CreateSessionRequest, user_session: str | None = Cookie(default=None)) -> dict:
    session_id = str(uuid.uuid4())
    with closing(get_connection()) as connection:
        user = require_user(connection, user_session)
        create_session_record(connection, session_id, user["id"], payload.model)
        session = get_session(connection, session_id, user["id"])
    return dict(session)


@app.get("/api/sessions/{session_id}")
def api_get_session(session_id: str, user_session: str | None = Cookie(default=None)) -> dict:
    with closing(get_connection()) as connection:
        user = require_user(connection, user_session)
        session = get_session(connection, session_id, user["id"])
        if session is None:
            raise HTTPException(status_code=404, detail="Session not found")
        return {
            "session": session,
            "messages": get_messages(connection, session_id, user["id"]),
        }


@app.patch("/api/sessions/{session_id}")
def api_update_session(
    session_id: str,
    payload: UpdateSessionRequest,
    user_session: str | None = Cookie(default=None),
) -> dict:
    with closing(get_connection()) as connection:
        user = require_user(connection, user_session)
        session = get_session(connection, session_id, user["id"])
        if session is None:
            raise HTTPException(status_code=404, detail="Session not found")
        update_session_metadata(connection, session_id, user["id"], session["model"], title=trim_title(payload.title))
        session = get_session(connection, session_id, user["id"])
    return {"session": session}


@app.delete("/api/sessions/{session_id}")
def api_delete_session(session_id: str, user_session: str | None = Cookie(default=None)) -> dict:
    with closing(get_connection()) as connection:
        user = require_user(connection, user_session)
        if not session_exists(connection, session_id, user["id"]):
            raise HTTPException(status_code=404, detail="Session not found")
        delete_session_record(connection, session_id, user["id"])
    return {"ok": True}


@app.delete("/api/sessions/{session_id}/messages/{message_id}")
def api_delete_message(session_id: str, message_id: int, user_session: str | None = Cookie(default=None)) -> dict:
    with closing(get_connection()) as connection:
        user = require_user(connection, user_session)
        session = get_session(connection, session_id, user["id"])
        if session is None:
            raise HTTPException(status_code=404, detail="Session not found")
        message = get_message(connection, session_id, user["id"], message_id)
        if message is None:
            raise HTTPException(status_code=404, detail="Message not found")

        delete_messages_from(connection, session_id, user["id"], message_id)
        remaining_messages = get_messages(connection, session_id, user["id"])

        if not remaining_messages:
            title = "New Chat"
        else:
            first_user = next((item for item in remaining_messages if item["role"] == "user"), None)
            title = (
                session["title"]
                if session["title"] != "New Chat"
                else build_title_for_message(
                    first_user["content"],
                    first_user.get("images") or [],
                    first_user.get("files") or [],
                ) if first_user else "New Chat"
            )

        update_session_metadata(connection, session_id, user["id"], session["model"], title=title)
        updated_session = get_session(connection, session_id, user["id"])

    return {
        "session": updated_session,
        "messages": remaining_messages,
    }


@app.post("/api/sessions/{session_id}/messages/{message_id}/resend")
def api_resend_message(
    session_id: str,
    message_id: int,
    payload: EditMessageRequest,
    user_session: str | None = Cookie(default=None),
) -> dict:
    message_text = payload.message.strip()
    if not message_text:
        raise HTTPException(status_code=400, detail="Message cannot be empty")

    with closing(get_connection()) as connection:
        user = require_user(connection, user_session)
        target_message = get_message(connection, session_id, user["id"], message_id)
        if target_message is None:
            raise HTTPException(status_code=404, detail="Message not found")
        if target_message["role"] != "user":
            raise HTTPException(status_code=400, detail="Only user messages can be edited and resent")
        if target_message.get("images") or target_message.get("files"):
            raise HTTPException(status_code=400, detail="Messages with attachments cannot be edited and resent yet")

        session, messages, assistant_text = send_and_persist_reply(
            connection=connection,
            session_id=session_id,
            user_id=user["id"],
            message_text=message_text,
            images=[],
            files=[],
            model=payload.model,
            replace_from_message_id=message_id,
        )

    return {
        "session": session,
        "messages": messages,
        "reply": assistant_text,
    }


@app.post("/api/chat")
def api_chat(payload: ChatRequest, user_session: str | None = Cookie(default=None)) -> dict:
    message_text = payload.message.strip()
    images = normalize_request_images(payload.image, payload.images)
    files = prepare_message_files(normalize_request_files(payload.files))
    if not message_text and not images and not files:
        raise HTTPException(status_code=400, detail="Message, image, or file is required")

    with closing(get_connection()) as connection:
        user = require_user(connection, user_session)
        session, messages, assistant_text = send_and_persist_reply(
            connection=connection,
            session_id=payload.session_id,
            user_id=user["id"],
            message_text=message_text,
            images=images,
            files=files,
            model=payload.model,
        )

    return {
        "session": session,
        "messages": messages,
        "reply": assistant_text,
    }


@app.post("/api/chat/stream")
def api_chat_stream(payload: ChatRequest, user_session: str | None = Cookie(default=None)) -> StreamingResponse:
    message_text = payload.message.strip()
    images = normalize_request_images(payload.image, payload.images)
    files = prepare_message_files(normalize_request_files(payload.files))
    if not message_text and not images and not files:
        raise HTTPException(status_code=400, detail="Message, image, or file is required")

    with closing(get_connection()) as connection:
        user = require_user(connection, user_session)
        removed_suffix: list[dict] = []
        replace_from_message_id = payload.replace_from_message_id

        if replace_from_message_id is not None:
            target_message = get_message(connection, payload.session_id, user["id"], replace_from_message_id)
            if target_message is None:
                raise HTTPException(status_code=404, detail="Message not found")
            if target_message["role"] != "user":
                raise HTTPException(status_code=400, detail="Only user messages can be edited and resent")
            if target_message.get("images") or target_message.get("files"):
                raise HTTPException(status_code=400, detail="Messages with attachments cannot be edited and resent yet")
            removed_suffix = [
                item
                for item in get_messages(
                    connection,
                    payload.session_id,
                    user["id"],
                    include_file_content=True,
                    include_original_data=True,
                )
                if item["id"] >= replace_from_message_id
            ]
            delete_messages_from(connection, payload.session_id, user["id"], replace_from_message_id)
        elif not session_exists(connection, payload.session_id, user["id"]):
            create_session_record(connection, payload.session_id, user["id"], payload.model)

        current_session = get_session(connection, payload.session_id, user["id"])
        history_before = get_messages(connection, payload.session_id, user["id"], include_file_content=True)
        add_message(connection, payload.session_id, "user", message_text, images=images, files=files)
        inserted_user = connection.execute(
            """
            SELECT id, session_id, role, content, created_at
            FROM messages
            WHERE session_id = ?
            ORDER BY id DESC
            LIMIT 1
            """,
            (payload.session_id,),
        ).fetchone()

        should_autobuild_title = not history_before and (current_session is None or current_session["title"] == "New Chat")
        title = build_title_for_message(message_text, images, files) if should_autobuild_title else None
        request_messages = build_request_messages(
            [*history_before, {"role": "user", "content": message_text, "images": images, "files": files}]
        )
        session = get_session(connection, payload.session_id, user["id"])
        retry_limit = max(1, min(get_api_key_count(connection), 50))

    def event_stream():
        assistant_parts: list[str] = []
        response_model = normalize_model(payload.model)

        try:
            yield sse_event(
                "start",
                {
                    "session": session,
                    "message": assign_message_files(
                        assign_message_images({"role": "user", "content": message_text}, images),
                        [
                            {
                                key: value
                                for key, value in file_payload.items()
                                if key not in {"data", "parsed_text"}
                            }
                            for file_payload in files
                        ],
                    ),
                },
            )

            attempted_key_ids: set[int] = set()
            last_exc: Exception | None = None
            stream_completed = False
            for _ in range(retry_limit):
                with closing(get_connection()) as key_connection:
                    key_record = ensure_user_api_key(key_connection, user["id"])

                key_id = int(key_record["id"])
                if key_id in attempted_key_ids:
                    break
                attempted_key_ids.add(key_id)

                emitted_text = False
                try:
                    for chunk in iter_stream_chat(
                        messages=request_messages,
                        model=payload.model,
                        session_id=payload.session_id,
                        api_key=key_record["api_key"],
                    ):
                        response_model = chunk.get("data", {}).get("model", response_model)
                        text_delta = extract_stream_text(chunk)
                        if text_delta:
                            emitted_text = True
                            assistant_parts.append(text_delta)
                            yield sse_event("delta", {"text": text_delta})
                    stream_completed = True
                    break
                except Exception as exc:
                    last_exc = exc
                    if emitted_text or not is_api_key_quota_error(exc):
                        raise
                    with closing(get_connection()) as rotate_connection:
                        advance_user_api_key(rotate_connection, user["id"], key_id)

            if not stream_completed:
                if last_exc is not None:
                    raise HTTPException(status_code=503, detail=f"所有可切换的 API Key 都不可用: {last_exc}") from last_exc
                raise HTTPException(status_code=503, detail="当前没有可用的 API Key")

            assistant_text = "".join(assistant_parts).strip() or "(empty response)"

            with closing(get_connection()) as connection:
                add_message(connection, payload.session_id, "assistant", assistant_text)
                update_session_metadata(connection, payload.session_id, user["id"], response_model, title=title)
                latest_session = get_session(connection, payload.session_id, user["id"])

            yield sse_event(
                "done",
                {
                    "session": latest_session,
                    "reply": assistant_text,
                },
            )
        except Exception as exc:
            with closing(get_connection()) as rollback_connection:
                if inserted_user is not None:
                    delete_message_files_by_ids(rollback_connection, [int(inserted_user["id"])])
                    rollback_connection.execute("DELETE FROM messages WHERE id = ?", (inserted_user["id"],))
                    rollback_connection.commit()
                if removed_suffix:
                    restore_messages(rollback_connection, removed_suffix)
            yield sse_event("error", {"detail": get_exception_detail(exc)})

    return StreamingResponse(event_stream(), media_type="text/event-stream")
